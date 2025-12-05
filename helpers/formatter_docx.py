# helpers/formatter_docx.py
import os
import re
import logging
from typing import Optional, Dict, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger("formatter_docx")
logger.setLevel(logging.INFO)


def _safe_run_boldify(paragraph, text: str):
    """
    Parsea **bold** dentro de text y añade runs al párrafo con bold donde toca.
    (Solo maneja pares de '**' sin anidamiento avanzado).
    """
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        m = re.match(r"^\*\*(.+)\*\*$", part)
        if m:
            run = paragraph.add_run(m.group(1))
            run.bold = True
        else:
            paragraph.add_run(part)


def insert_header_images(doc: Document, banner_path: Optional[str], logo_path: Optional[str], banner_height_in=Inches(1.0), logo_width_in=Inches(1.0)):
    """
    Inserta banner (izq/centro superior) y logo (derecha superior) en el header de la primera sección.
    Si no existen rutas, no hace nada.
    """
    try:
        section = doc.sections[0]
        header = section.header
        # Clear existing header paragraphs for predictable layout
        for p in header.paragraphs:
            p.clear()

        # Create a table in header: 1 row x 2 cols -> left banner, right logo
        tbl = header.add_table(rows=1, cols=2)
        tbl.autofit = False
        left_cell = tbl.cell(0, 0)
        right_cell = tbl.cell(0, 1)

        # set widths if possible (best-effort)
        try:
            left_cell.width = Inches(6.0)
            right_cell.width = Inches(1.5)
        except Exception:
            pass

        if banner_path and os.path.exists(banner_path):
            p = left_cell.paragraphs[0]
            run = p.add_run()
            try:
                run.add_picture(banner_path, width=banner_height_in * 8)  # best-effort wide
            except Exception:
                logger.exception("No se pudo insertar banner en header: %s", banner_path)

        if logo_path and os.path.exists(logo_path):
            p = right_cell.paragraphs[0]
            p.alignment = 2  # right
            run = p.add_run()
            try:
                run.add_picture(logo_path, width=logo_width_in)
            except Exception:
                logger.exception("No se pudo insertar logo en header: %s", logo_path)

    except Exception as e:
        logger.exception("insert_header_images error: %s", e)


def add_colored_heading(doc: Document, text: str, color_rgb=(3, 94, 99)):
    """
    Crea una tabla 1x1 con shading (barra de color) y texto blanco en su interior.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = 0
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run(text.strip())
    run.bold = True
    # white font
    try:
        run.font.color.rgb = RGBColor(255, 255, 255)
    except Exception:
        pass

    # shading
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), "%02x%02x%02x" % color_rgb)
    tcPr.append(sh)


def _add_paragraph_with_style(doc: Document, text: str, style_name: Optional[str] = None):
    if style_name:
        p = doc.add_paragraph(style=style_name)
        _safe_run_boldify(p, text)
    else:
        p = doc.add_paragraph()
        _safe_run_boldify(p, text)
    return p


def _render_markdown_to_docx(doc: Document, markdown_text: str, use_colored_bar: bool = True, bar_color=(3, 94, 99)):
    """
    Simple markdown -> docx renderer (supports ##, ###, - list, paragraphs and **bold**).
    Not a full markdown engine, but sufficient for our generated TCPs.
    """
    lines = markdown_text.splitlines()
    i = 0
    list_buffer: List[str] = []

    def flush_list():
        nonlocal list_buffer
        if not list_buffer:
            return
        for li in list_buffer:
            p = doc.add_paragraph(style='List Bullet')
            _safe_run_boldify(p, li.strip()[2:].strip())
        list_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()
        i += 1
        if not line:
            flush_list()
            continue

        # Heading level 2 (##)
        if line.startswith("## "):
            flush_list()
            heading_text = line[3:].strip()
            # Decide whether to render as colored bar or normal heading
            if use_colored_bar:
                # if style 'Reda_Section' exists in template, prefer simple heading with that style
                if any(s.name == 'Reda_Section' for s in doc.styles):
                    p = doc.add_paragraph(style='Reda_Section')
                    _safe_run_boldify(p, heading_text)
                else:
                    add_colored_heading(doc, heading_text, color_rgb=bar_color)
            else:
                # Use Reda_Title or Heading 1
                style_name = 'Reda_Title' if any(s.name == 'Reda_Title' for s in doc.styles) else 'Heading 1'
                _add_paragraph_with_style(doc, heading_text, style_name)
            continue

        # Heading level 3 (###)
        if line.startswith("### "):
            flush_list()
            heading_text = line[4:].strip()
            style_name = 'Reda_Subtitle' if any(s.name == 'Reda_Subtitle' for s in doc.styles) else 'Heading 2'
            _add_paragraph_with_style(doc, heading_text, style_name)
            continue

        # Unordered list
        if line.startswith("- "):
            list_buffer.append(line)
            # keep collecting subsequent list items
            while i < len(lines) and lines[i].lstrip().startswith("- "):
                list_buffer.append(lines[i].rstrip())
                i += 1
            flush_list()
            continue

        # Regular paragraph
        flush_list()
        p = doc.add_paragraph()
        _safe_run_boldify(p, line)

    flush_list()


def replace_marker_in_docx(template_path: str, output_path: str, content_markdown: str,
                          banner_path: Optional[str] = None, logo_path: Optional[str] = None,
                          quiz_text: Optional[str] = None, images_map: Optional[Dict[int, List[str]]] = None,
                          use_colored_bar: bool = True):
    """
    - template_path: plantilla .docx con un marcador <!--REDA_CONTENT--> dentro del body (puede estar en un paragraph).
    - output_path: ruta donde guardar el docx final.
    - content_markdown: contenido (Markdown) a insertar.
    - banner_path / logo_path: rutas opcionales a imágenes (se buscan usualmente en templates/images/).
    - quiz_text: si lo pasas, se inserta en una nueva página al final con título "RedaQuiz".
    - images_map: dict page_number -> list of image paths (best-effort: will append images at the start of each section).
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    # Insert header images (best-effort)
    try:
        insert_header_images(doc, banner_path, logo_path)
    except Exception:
        logger.exception("Failed to insert header images - continuing")

    # Find marker paragraph(s)
    marker = "<!--REDA_CONTENT-->"
    replaced = False

    for para in list(doc.paragraphs):
        if marker in para.text:
            # Clear paragraph runs
            p_parent = para._p
            # remove runs
            for r in list(para.runs):
                r.clear()
            # replace the paragraph by inserting content here: we'll insert a temporary paragraph index
            # docx API doesn't give exact insertion before para easily; we'll append at end and later reorder if needed.
            # Simpler approach: remove the marker paragraph text and then insert content immediately after it.
            para.text = ""
            # Insert content right after this paragraph by finding its index
            # We'll build new doc by iterating paragraphs and write to a new Document
            replaced = True
            break

    if not replaced:
        logger.warning("Marker %s not found in template. Appending content at the end.", marker)

    # Approach: create a new document, copy header/footer/sections from template, then append content.
    # Simpler: keep same doc object and append content at the marker location (we cleared it above).
    # Append content now:
    _render_markdown_to_docx(doc, content_markdown, use_colored_bar=use_colored_bar)

    # Insert images_map if provided: best-effort append images in doc (grouped by page/section)
    if images_map:
        try:
            doc.add_page_break()
            for page_idx, imgs in images_map.items():
                doc.add_paragraph(f"Images for page {page_idx}:")
                for img in imgs:
                    if os.path.exists(img):
                        try:
                            r = doc.add_paragraph().add_run()
                            r.add_picture(img, width=Inches(4.5))
                        except Exception:
                            logger.exception("Failed inserting image %s", img)
                    else:
                        logger.warning("Image path not found in images_map: %s", img)
        except Exception:
            logger.exception("Failed to insert images_map")

    # Insert RedaQuiz if provided (in new page)
    if quiz_text:
        try:
            doc.add_page_break()
            style_name = 'Reda_Section' if any(s.name == 'Reda_Section' for s in doc.styles) else 'Heading 1'
            _add_paragraph_with_style(doc, "RedaQuiz", style_name)
            # quiz as plain preformatted text paragraphs (preserve lines)
            for line in quiz_text.splitlines():
                doc.add_paragraph(line)
        except Exception:
            logger.exception("Failed inserting RedaQuiz")

    # Footer: ensure simple footer with RedaXion if none
    try:
        section = doc.sections[0]
        footer = section.footer
        if not footer.paragraphs or not any(p.text.strip() for p in footer.paragraphs):
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.text = "RedaXion"
            p.alignment = 1  # center
    except Exception:
        logger.exception("Failed ensure footer")

    # Save document
    try:
        doc.save(output_path)
        logger.info("Saved formatted docx to %s", output_path)
    except Exception:
        logger.exception("Failed saving document to %s", output_path)
        raise


# Small CLI-style helper for manual testing (local)
if __name__ == "__main__":
    import argparse
    p = argparse.ArgumentParser(description="Apply template and insert Markdown content into DOCX")
    p.add_argument("--template", required=True, help="Path to template .docx")
    p.add_argument("--out", required=True, help="Output .docx path")
    p.add_argument("--md", required=True, help="Path to markdown/text file to insert")
    p.add_argument("--banner", required=False, help="Path to banner image (optional)")
    p.add_argument("--logo", required=False, help="Path to logo image (optional)")
    p.add_argument("--quiz", required=False, help="Path to quiz txt (optional)")
    args = p.parse_args()

    md_text = ""
    with open(args.md, "r", encoding="utf-8") as fh:
        md_text = fh.read()

    quiz_text = None
    if args.quiz and os.path.exists(args.quiz):
        with open(args.quiz, "r", encoding="utf-8") as fh:
            quiz_text = fh.read()

    replace_marker_in_docx(
        template_path=args.template,
        output_path=args.out,
        content_markdown=md_text,
        banner_path=args.banner,
        logo_path=args.logo,
        quiz_text=quiz_text,
        images_map=None,
        use_colored_bar=True,
    )
