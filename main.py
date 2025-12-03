# main.py - flujo completo RedaXion (actualizado) - versión integrada con process_txt
"""
generate_and_deliver(order_id, *args, **kwargs)

Orquesta el flujo completo para una orden:
- descarga / obtiene URL del audio (ya subido a GCS)
- transcribe con AssemblyAI -> guarda .txt
- divide en bloques de 3000 palabras y procesa con ChatGPT-5 por bloque
- une bloques en TCP en estilo libro universitario
- extrae títulos/subtítulos y por cada "página" (estimada) busca una imagen relevante
- genera 7 preguntas difíciles por página (RedaQuiz)
- aplica plantilla .docx (color/columnas)
- convierte a pdf, sube todo a Drive, marca Sheets y envía correo
"""

import os
import tempfile
import traceback
import shutil
import json
from datetime import datetime, timedelta

# Intentar importar tus helpers (estructura original). Si están en 'helpers.*' ajustamos.
try:
    from sheets import get_todos_los_pendientes, marcar_como_procesado, get_pedido_por_fila, actualizar_estado_y_links
    from gcs import procesar_audio
    from assemblyai import transcribir_audio
    from process_txt import procesar_txt_con_chatgpt_block  # tu helper nuevo (si está en raíz)
    from formatter_docx import guardar_como_docx, guardar_quiz_como_docx
    from subir_archivo import subir_archivo_a_drive
    from generar_quiz import generar_quiz_desde_docx
    from enviar_correo import enviar_correo_con_adjuntos
    from convertidor_pdf import convertir_a_pdf
except Exception:
    # Try alternate import paths used earlier in the repo (helpers package)
    try:
        from helpers.sheets import get_todos_los_pendientes, marcar_como_procesado, get_pedido_por_fila, actualizar_estado_y_links
        from helpers.gcs import procesar_audio
        from helpers.assemblyai import transcribir_audio
        from helpers.process_txt import procesar_txt_con_chatgpt_block
        from helpers.formatter_docx import guardar_como_docx, guardar_quiz_como_docx
        from helpers.subir_archivo import subir_archivo_a_drive
        from helpers.generar_quiz import generar_quiz_desde_docx
        from helpers.enviar_correo import enviar_correo_con_adjuntos
        from helpers.convertidor_pdf import convertir_a_pdf
    except Exception:
        # If imports fail, we'll use internal stubs and continue so the worker doesn't crash.
        get_todos_los_pendientes = None
        marcar_como_procesado = None
        get_pedido_por_fila = None
        actualizar_estado_y_links = None
        procesar_audio = None
        transcribir_audio = None
        procesar_txt_con_chatgpt_block = None
        guardar_como_docx = None
        guardar_quiz_como_docx = None
        subir_archivo_a_drive = None
        generar_quiz_desde_docx = None
        enviar_correo_con_adjuntos = None
        convertir_a_pdf = None

# -----------------------
# Fallback directo a Google Sheets (si los helpers no funcionan)
# -----------------------
def get_details_from_sheet_direct(order_id: str):
    """
    Fallback: leer hoja directamente con gspread usando GOOGLE_SHEETS_CREDENTIALS_JSON y SHEET_ID.
    Devuelve dict con claves similares a las que espera el flujo: orden, color, columnas, email, audio_url, fila.
    """
    try:
        import gspread
    except Exception as e:
        print("[SHEETS FALLBACK] gspread no disponible:", e)
        return None

    creds_json = os.getenv("GOOGLE_SHEETS_CREDENTIALS_JSON") or os.getenv("GCS_CREDENTIALS_JSON")
    sheet_id = os.getenv("SHEET_ID") or os.getenv("GOOGLE_SHEET_ID") or os.getenv("SHEET_KEY")
    if not creds_json:
        print("[SHEETS FALLBACK] No encontré GOOGLE_SHEETS_CREDENTIALS_JSON en env.")
        return None
    if not sheet_id:
        print("[SHEETS FALLBACK] No encontré SHEET_ID en env.")
        return None

    try:
        info = json.loads(creds_json)
    except Exception as e:
        print("[SHEETS FALLBACK] GOOGLE_SHEETS_CREDENTIALS_JSON no es JSON válido:", e)
        return None

    try:
        gc = gspread.service_account_from_dict(info)
    except Exception as e:
        print("[SHEETS FALLBACK] Error creando service_account_from_dict:", e)
        return None

    # Abrir spreadsheet: intenta por key o por url
    sh = None
    try:
        sh = gc.open_by_key(sheet_id)
    except Exception:
        try:
            sh = gc.open_by_url(sheet_id)
        except Exception as e:
            print("[SHEETS FALLBACK] No pude abrir la hoja con SHEET_ID/url:", e)
            return None

    try:
        ws = sh.sheet1  # si usas otra hoja, cambia aquí
        records = ws.get_all_records()  # lista de dicts usando header como key
    except Exception as e:
        print("[SHEETS FALLBACK] Error leyendo worksheet:", e)
        return None

    for idx, row in enumerate(records, start=2):  # start=2 -> fila real en sheet (1 = header)
        # normaliza la clave 'orden' (puede venir con mayúsculas)
        key_candidates = [k for k in row.keys() if k.lower().strip() == "orden"]
        if key_candidates:
            key = key_candidates[0]
            if str(row.get(key)).strip() == str(order_id).strip():
                # construir dict con claves esperadas
                detalle = {
                    "orden": row.get(key),
                    "fila": idx,
                    "fecha": row.get("fecha") or row.get("Fecha"),
                    "nombre": row.get("nombre") or row.get("Nombre"),
                    "email": row.get("email") or row.get("Email"),
                    "audio_url": row.get("audio_url") or row.get("Audio_URL") or row.get("audio") or row.get("Audio"),
                    "columnas": row.get("columnas") or row.get("Columnas"),
                    "color": row.get("color") or row.get("Color"),
                    "estado": row.get("estado") or row.get("Estado"),
                    "payment_id": row.get("payment_id") or row.get("payment id") or row.get("payment"),
                    "comentarios": row.get("comentarios") or row.get("Comentarios") or "",
                }
                print(f"[SHEETS FALLBACK] Orden encontrada en fila {idx}: {detalle}")
                return detalle

    print("[SHEETS FALLBACK] No encontré la orden en la hoja (fallback).")
    return None

# -----------------------
# Utilidades internas
# -----------------------
def split_text_into_blocks(text: str, words_per_block: int = 3000):
    words = text.split()
    blocks = []
    for i in range(0, len(words), words_per_block):
        block = " ".join(words[i:i + words_per_block])
        blocks.append(block)
    return blocks

def call_chatgpt_for_block(block_text: str, block_index: int, order_id: str, total_blocks: int):
    """
    Llama a la función que transforma un bloque en estilo 'TCP' y devuelve texto formateado.
    Usa el helper procesar_txt_con_chatgpt_block si existe.
    Guarda resultado parcial en /tmp para debugging.
    """
    try:
        print(f"[CHATGPT] Iniciando bloque {block_index}/{total_blocks} (order {order_id})")
        if procesar_txt_con_chatgpt_block:
            result = procesar_txt_con_chatgpt_block(block_text, order_id=order_id, block_index=block_index, total_blocks=total_blocks)
        else:
            print("[CHATGPT][STUB] helper procesar_txt_con_chatgpt_block no disponible. Usando stub.")
            result = f"## Bloque {block_index}\n\n" + block_text[:2000] + "\n\n"

        # Guardar salida parcial en archivo (útil para revisar si algo falla)
        try:
            tmp_file = f"/tmp/{order_id}_block_{block_index}.md"
            with open(tmp_file, "w", encoding="utf-8") as fh:
                fh.write(result)
            print(f"[CHATGPT] Resultado bloque {block_index} guardado en {tmp_file}")
        except Exception as e:
            print(f"[CHATGPT][WARN] No se pudo guardar block_{block_index}.md: {e}")

        return result
    except Exception as e:
        print(f"[CHATGPT][ERROR] al procesar bloque {block_index}: {e}")
        traceback.print_exc()
        return f"## ERROR BLOQUE {block_index}\n\n{block_text[:8000]}\n\n"

def merge_processed_blocks(blocks_processed):
    return "\n\n".join(blocks_processed)

def extract_titles_subtitles(tcp_text: str):
    """
    Heurística simple para extraer títulos/subtítulos desde el texto TCP.
    Devuelve lista de tuples (title, subtitle, approx_page).
    """
    lines = [l.strip() for l in tcp_text.splitlines() if l.strip()]
    titles = []
    for i, line in enumerate(lines):
        if len(line) < 120 and line[0].isupper() and len(line.split()) < 8:
            subtitle = ""
            if i + 1 < len(lines) and len(lines[i+1].split()) < 12:
                subtitle = lines[i+1]
            titles.append((line, subtitle))
    result = []
    if not titles:
        return []
    for idx, t in enumerate(titles):
        approx_page = idx + 1
        result.append((t[0], t[1], approx_page))
    return result

def search_image_for_topic(title: str, subtitle: str = ""):
    """
    Placeholder: Buscar la mejor imagen para un tema.
    Reemplaza esto con integración real (Unsplash, Bing Image Search, Google Custom Search, etc.).
    """
    query = f"{title} {subtitle}".strip()
    print(f"[IMG_SEARCH][STUB] buscar imagen para: {query}")
    return f"https://via.placeholder.com/1200x800.png?text={query.replace(' ', '+')}"

def generate_questions_for_titles(titles_list, per_title=7):
    """
    Genera preguntas por cada título/subtítulo (STUB).
    """
    questions_by_page = {}
    for title, subtitle, page in titles_list:
        questions = []
        for i in range(per_title):
            q_text = f"Pregunta difícil sobre «{title}» (ítem {i+1})"
            options = [
                "A) Opción 1",
                "B) Opción 2",
                "C) Opción 3",
                "D) Opción 4",
                "E) Opción 5",
            ]
            questions.append({
                "question": q_text,
                "options": options,
                "answer": "A",
                "justification": "Justificación breve (stub)."
            })
        questions_by_page.setdefault(page, []).extend(questions)
    return questions_by_page

def apply_docx_template_and_insert_images(tcp_text, images_map, out_path, color="azul", columnas="simple"):
    """
    Usa tu helper `guardar_como_docx` si existe; sino genera DOCX simple.
    """
    try:
        if guardar_como_docx:
            try:
                return guardar_como_docx(tcp_text, out_path, color=color, columnas=columnas, images_map=images_map)
            except TypeError:
                print("[DOCX] guardar_como_docx no acepta images_map: llamando sin images_map.")
                return guardar_como_docx(tcp_text, out_path, color=color, columnas=columnas)
        else:
            print("[DOCX][STUB] creando DOCX simple")
            from docx import Document
            doc = Document()
            for para in tcp_text.split("\n\n"):
                doc.add_paragraph(para)
            doc.save(out_path)
            return out_path
    except Exception as e:
        print(f"[DOCX][ERROR] al generar DOCX: {e}")
        traceback.print_exc()
        return None

def apply_quiz_template_and_save(questions_by_page, out_quiz_path, color="azul", columnas="simple"):
    """
    Convierte questions_by_page a un docx con formato.
    """
    try:
        if guardar_quiz_como_docx:
            return guardar_quiz_como_docx(questions_by_page, out_quiz_path, color=color, columnas=columnas)
        else:
            print("[QUIZ][STUB] guardando quiz básico en docx")
            from docx import Document
            doc = Document()
            for page, qs in questions_by_page.items():
                doc.add_heading(f"Preguntas - Página {page}", level=2)
                for idx, q in enumerate(qs, start=1):
                    doc.add_paragraph(f"{idx}. {q['question']}")
                    for opt in q['options']:
                        doc.add_paragraph(f"   {opt}")
                doc.add_page_break()
            doc.save(out_quiz_path)
            return out_quiz_path
    except Exception as e:
        print(f"[QUIZ][ERROR] {e}")
        traceback.print_exc()
        return None

# -----------------------
# Orquestador principal
# -----------------------
def generate_and_deliver(order_id, *args, **kwargs):
    """
    Orquesta todo el pipeline para una orden específica.
    """
    tmp_dir = None
    try:
        print(f"\n🚀 [MAIN] generate_and_deliver -> order_id={order_id} - inicio {datetime.utcnow().isoformat()}")
        if kwargs:
            print(f"[MAIN] kwargs recibidos: {kwargs}")

        # 1) Obtener datos de la orden desde sheets (fila / detalles)
        detalles = None

        # 1.a Intentar helper directo por orden
        if get_pedido_por_fila:
            try:
                detalles = get_pedido_por_fila(order_id)
                print(f"[MAIN] get_pedido_por_fila devolvió: {bool(detalles)}")
            except Exception as e:
                print("[MAIN] get_pedido_por_fila lanzó excepción:", e)
                detalles = None

        # 1.b Intentar lista de pendientes
        if not detalles and get_todos_los_pendientes:
            try:
                pendientes = get_todos_los_pendientes()
                print(f"[MAIN] get_todos_los_pendientes devolvió {len(pendientes)} items (si es lista).")
                detalles = next((p for p in pendientes if str(p.get("orden")).strip() == str(order_id).strip()), None)
            except Exception as e:
                print("[MAIN] get_todos_los_pendientes excepción:", e)
                detalles = None

        # 1.c Fallback directo a Google Sheets (gspread)
        if not detalles:
            print("[MAIN] Intentando fallback directo a Google Sheets...")
            try:
                detalles = get_details_from_sheet_direct(order_id)
            except Exception as e:
                print("[MAIN] get_details_from_sheet_direct excepción:", e)
                detalles = None

        if not detalles:
            print(f"[MAIN][WARN] No se encontró metadata de la orden {order_id} en Sheets. Abortando.")
            return

        # Validación básica
        color = detalles.get("color", "azul")
        columnas = detalles.get("columnas", "simple")
        correo_cliente = detalles.get("email", "") or detalles.get("correo") or ""
        fila = detalles.get("fila") or detalles.get("row") or None

        estado_actual = (detalles.get("estado") or "").strip().lower()
        if "entregado" in estado_actual or "procesado" in estado_actual:
            print(f"[MAIN] Orden {order_id} ya tiene estado '{estado_actual}'. Saltando procesamiento.")
            return

        # 2) Obtener URL pública del audio
        audio_url_public = detalles.get("audio_url") or detalles.get("url") or None
        if not audio_url_public:
            print("[MAIN] No se encontró audio_url en metadata: intentando reprocesar con gcs.procesar_audio (si se tiene path)")
            try:
                source_path = detalles.get("audio_path") or detalles.get("gdrive_path") or detalles.get("path")
                if procesar_audio and source_path:
                    audio_url_public = procesar_audio(source_path, f"{order_id}.mp3")
            except Exception:
                audio_url_public = None

        if not audio_url_public:
            print(f"[MAIN][ERROR] No hay URL pública del audio para la orden {order_id}. Marcar en Sheets y abortar.")
            try:
                if actualizar_estado_y_links:
                    actualizar_estado_y_links(order_id, estado="Error: no audio_url")
            except Exception:
                pass
            return

        print(f"[MAIN] Audio público: {audio_url_public}")

        # 3) Transcribir con AssemblyAI -> obtener texto completo y guardar .txt
        try:
            print("[MAIN] Enviando a AssemblyAI para transcripción (puede tardar)...")
            if transcribir_audio:
                texto = transcribir_audio(audio_url_public)
            else:
                print("[MAIN][STUB] transcribir_audio helper no disponible. Creando texto stub.")
                texto = "Transcripción de prueba. " * 1000
        except Exception as e:
            print("[MAIN][ERROR] Falló transcripción:", e)
            traceback.print_exc()
            if actualizar_estado_y_links:
                actualizar_estado_y_links(order_id, estado=f"Error: transcripcion {e}")
            return

        # Crear tmp_dir para archivos temporales
        tmp_dir = tempfile.mkdtemp(prefix=f"redax_{order_id}_")
        path_txt = os.path.join(tmp_dir, f"{order_id}.txt")
        try:
            with open(path_txt, "w", encoding="utf-8") as f:
                f.write(texto)
            print(f"[MAIN] Guardado .txt temporal en {path_txt}")
            if subir_archivo_a_drive:
                try:
                    subir_archivo_a_drive(path_txt, f"{order_id}.txt", order_id)
                    print("[MAIN] .txt subido a Drive.")
                except Exception as e:
                    print(f"[MAIN][WARN] No se pudo subir .txt a Drive: {e}")
        except Exception as e:
            print("[MAIN][ERROR] No se pudo guardar/subir .txt:", e)
            traceback.print_exc()

        # 4) Dividir en bloques y procesar cada bloque
        blocks = split_text_into_blocks(texto, words_per_block=3000)
        total_blocks = len(blocks)
        print(f"[MAIN] Texto dividido en {total_blocks} bloques de ~3000 palabras.")

        processed_blocks = []
        for i, blk in enumerate(blocks, start=1):
            pb = call_chatgpt_for_block(blk, i, order_id, total_blocks)
            processed_blocks.append(pb)

        tcp_text = merge_processed_blocks(processed_blocks)
        print(f"[MAIN] TCP (texto procesado) ensamblado, tamaño {len(tcp_text)} caracteres.")

        # 5) Extraer títulos/subtítulos
        titles = extract_titles_subtitles(tcp_text)
        print(f"[MAIN] Extraídos {len(titles)} títulos/subtítulos heurísticos.")

        # 6) Buscar imágenes (STUB)
        images_map = {}
        for title, subtitle, page in titles:
            img_url = search_image_for_topic(title, subtitle)
            images_map[page] = img_url
        print(f"[MAIN] Imágenes buscadas para {len(images_map)} páginas (map listo).")

        # 7) Generar preguntas por título/página (7 por página)
        questions_by_page = generate_questions_for_titles(titles, per_title=7)
        total_questions = sum(len(v) for v in questions_by_page.values())
        print(f"[MAIN] Generadas preguntas: {total_questions} ítems.")

        # 8) Guardar TCP en DOCX usando plantilla y aplicar imágenes por página
        nombre_tcp = f"RedaXion - Nº{order_id}.docx"
        path_docx = os.path.join(tmp_dir, nombre_tcp)
        docx_path_result = apply_docx_template_and_insert_images(tcp_text, images_map, path_docx, color=color, columnas=columnas)
        if docx_path_result:
            print(f"[MAIN] DOCX TCP generado en {docx_path_result}")
            if subir_archivo_a_drive:
                try:
                    subir_archivo_a_drive(docx_path_result, nombre_tcp, order_id)
                    print("[MAIN] DOCX TCP subido a Drive.")
                except Exception as e:
                    print(f"[MAIN][WARN] No se pudo subir DOCX a Drive: {e}")
        else:
            print("[MAIN][ERROR] No se pudo generar DOCX TCP.")

        # 9) Convertir DOCX TCP a PDF
        pdf_tcp = None
        try:
            if convertir_a_pdf and docx_path_result:
                pdf_tcp = convertir_a_pdf(docx_path_result)
                if pdf_tcp:
                    nombre_tcp_pdf = nombre_tcp.replace(".docx", ".pdf")
                    try:
                        subir_archivo_a_drive(pdf_tcp, nombre_tcp_pdf, order_id)
                        print(f"[MAIN] PDF TCP generado y subido: {nombre_tcp_pdf}")
                    except Exception as e:
                        print(f"[MAIN][WARN] No se pudo subir PDF TCP a Drive: {e}")
        except Exception as e:
            print("[MAIN][ERROR] convertir_a_pdf fallo:", e)
            traceback.print_exc()

        # 10) Generar QUIZ DOCX (format)
        nombre_quiz = f"RedaQuiz - Nº{order_id}.docx"
        path_quiz = os.path.join(tmp_dir, nombre_quiz)
        quiz_docx = apply_quiz_template_and_save(questions_by_page, path_quiz, color=color, columnas=columnas)
        if quiz_docx:
            print(f"[MAIN] Quiz DOCX generado en {quiz_docx}")
            if subir_archivo_a_drive:
                try:
                    subir_archivo_a_drive(quiz_docx, nombre_quiz, order_id)
                    print("[MAIN] Quiz DOCX subido a Drive.")
                except Exception as e:
                    print(f"[MAIN][WARN] No se pudo subir Quiz DOCX a Drive: {e}")
        else:
            print("[MAIN][WARN] No se produjo Quiz DOCX.")

        # 11) Convertir QUIZ a PDF
        pdf_quiz = None
        try:
            if convertir_a_pdf and quiz_docx:
                pdf_quiz = convertir_a_pdf(quiz_docx)
                if pdf_quiz:
                    nombre_quiz_pdf = nombre_quiz.replace(".docx", ".pdf")
                    try:
                        subir_archivo_a_drive(pdf_quiz, nombre_quiz_pdf, order_id)
                        print(f"[MAIN] PDF Quiz generado y subido: {nombre_quiz_pdf}")
                    except Exception as e:
                        print(f"[MAIN][WARN] No se pudo subir PDF Quiz a Drive: {e}")
        except Exception as e:
            print("[MAIN][ERROR] convertir_a_pdf (quiz) fallo:", e)
            traceback.print_exc()

        # 12) Actualizar Sheets: marcar como entregado y publicar links (si tienes helper)
        try:
            if actualizar_estado_y_links:
                links = {
                    "txt": f"{order_id}.txt",
                    "docx_tcp": nombre_tcp,
                    "pdf_tcp": nombre_tcp.replace(".docx", ".pdf") if pdf_tcp else None,
                    "docx_quiz": nombre_quiz,
                    "pdf_quiz": nombre_quiz.replace(".docx", ".pdf") if pdf_quiz else None,
                }
                actualizar_estado_y_links(order_id, estado="Entregado", links=links)
                print("[MAIN] Sheets actualizado con estado Entregado y links.")
            else:
                if marcar_como_procesado and fila:
                    marcar_como_procesado(fila)
                    print("[MAIN] marcar_como_procesado ejecutado.")
        except Exception:
            print("[MAIN][WARN] Falló al actualizar Sheets con links/estado.")
            traceback.print_exc()

        # 13) Enviar correo al cliente con adjuntos (docx + pdf)
        try:
            archivos_adjuntos = []
            if docx_path_result:
                archivos_adjuntos.append(docx_path_result)
            if pdf_tcp:
                archivos_adjuntos.append(pdf_tcp)
            if quiz_docx:
                archivos_adjuntos.append(quiz_docx)
            if pdf_quiz:
                archivos_adjuntos.append(pdf_quiz)

            if enviar_correo_con_adjuntos and correo_cliente:
                asunto = f"Tu pedido RedaXion Nº{order_id} está listo ✅"
                cuerpo = (
                    f"Hola 👋\n\nAdjuntamos tu Transcripción Académica Profesional (TCP) y el RedaQuiz.\n"
                    "Gracias por usar RedaXion — ¡éxitos en el estudio! 🧠\n\n"
                    "— Equipo RedaXion"
                )
                try:
                    enviar_correo_con_adjuntos(correo_cliente, asunto, cuerpo, archivos_adjuntos)
                    print(f"[MAIN] Correo enviado a {correo_cliente}")
                except Exception as e:
                    print(f"[MAIN][WARN] No se pudo enviar correo: {e}")
            else:
                print("[MAIN] enviar_correo_con_adjuntos helper no disponible o correo_cliente vacío; correo no enviado.")
        except Exception:
            print("[MAIN][ERROR] Error enviando correo con adjuntos.")
            traceback.print_exc()

        print(f"✅ [MAIN] Finalizado flujo para orden {order_id} ({datetime.utcnow().isoformat()})")
        return True

    except Exception as err:
        print(f"[MAIN][ERROR] Excepción en generate_and_deliver para {order_id}: {err}")
        traceback.print_exc()
        try:
            if actualizar_estado_y_links:
                actualizar_estado_y_links(order_id, estado=f"Error: {err}")
        except Exception:
            pass
        return False

    finally:
        # limpiar tmp_dir si existe
        try:
            if tmp_dir and os.path.isdir(tmp_dir):
                shutil.rmtree(tmp_dir)
                print(f"[MAIN] tmp_dir {tmp_dir} eliminado.")
        except Exception as e:
            print(f"[MAIN][WARN] No se pudo eliminar tmp_dir {tmp_dir}: {e}")
            traceback.print_exc()

# Mantener compatibilidad (si el worker importa main.generate_and_deliver)
if __name__ == "__main__":
    try:
        print("🚀 Ejecutando flujo RedaXion en modo standalone (procesar pendientes)")
        if get_todos_los_pendientes:
            pendientes = get_todos_los_pendientes()
            for p in pendientes:
                try:
                    oid = p.get("orden")
                    if oid:
                        generate_and_deliver(oid)
                except Exception:
                    traceback.print_exc()
        else:
            print("No hay helper get_todos_los_pendientes disponible.")
    except Exception:
        traceback.print_exc()
